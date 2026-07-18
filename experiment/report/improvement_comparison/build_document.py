#!/usr/bin/env python3
"""Build the standalone improvement-to-paper comparison report.

The document intentionally separates protocol-aligned comparisons from
engineering ensemble points and internal 400-row experiments.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# The script lives at <repo>/experiment/report/improvement_comparison.
# Deriving the root keeps regeneration independent of the current directory.
ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "experiment" / "report" / "improvement_comparison"
RESULTS = ROOT / "experiment" / "results"
FIGURES = ROOT / "experiment" / "figures"
PAPER_ASSETS = ROOT / "experiment" / "paper_reader" / "assets"

DOCX_PATH = OUT / "提升数据与原论文逐项对比.docx"
MD_PATH = OUT / "提升数据与原论文逐项对比.md"
TABLE5_CROP = OUT / "qa" / "table5_recdy_crop.png"

NAVY = "17324D"
TEAL = "237C82"
TEAL_LIGHT = "DCEEEF"
BLUE_LIGHT = "E9F0F7"
GOLD = "B07A16"
GOLD_LIGHT = "F8EED7"
RED = "A84A59"
RED_LIGHT = "F7E5E8"
GREEN = "2D7C55"
GREEN_LIGHT = "E3F1EA"
GRAY = "5F6872"
GRAY_LIGHT = "F1F3F5"
WHITE = "FFFFFF"


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def percent(value: float, digits: int = 2) -> str:
    return f"{100 * value:.{digits}f}%"


def pp(value: float, digits: int = 2) -> str:
    return f"{value:+.{digits}f} pp"


def soft_wrap_path(value: str, max_chars: int = 48) -> str:
    """Wrap a Windows path at directory separators without changing it."""
    parts = value.split("\\")
    lines: list[str] = []
    current = parts[0]
    for part in parts[1:]:
        candidate = f"{current}\\{part}" if current else part
        if current and len(candidate) > max_chars:
            lines.append(f"{current}\\")
            current = part
        else:
            current = candidate
    if current:
        lines.append(current)
    return "\n".join(lines)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edges[edge].items():
            element.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=75, start=85, bottom=75, end=85) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=8, color=GRAY)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=17, color=NAVY, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=TEAL, bold=True)
    else:
        set_run_font(run, size=11, color=NAVY, bold=True)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=7, line=1.35)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=9.5, color=NAVY, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=9.5, color="27313B")
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.5, color="27313B")


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph_spacing(p, after=4, line=1.25)
    run = p.add_run(text)
    set_run_font(run, size=9.2, color="27313B")


def add_callout(doc: Document, label: str, text: str, kind: str = "info") -> None:
    palette = {
        "info": (TEAL, TEAL_LIGHT),
        "warning": (GOLD, GOLD_LIGHT),
        "risk": (RED, RED_LIGHT),
        "success": (GREEN, GREEN_LIGHT),
    }
    accent, fill = palette[kind]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.45)
    table.columns[1].width = Cm(14.3)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=120, bottom=120)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 4, "color": accent},
            bottom={"val": "single", "sz": 4, "color": accent},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left.text = ""
    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(label)
    set_run_font(r1, size=8.5, color=accent, bold=True)
    right.text = ""
    p2 = right.paragraphs[0]
    r2 = p2.add_run(text)
    set_run_font(r2, size=8.8, color="27313B")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None,
              header_fill=NAVY, highlight_rows=None, font_size=7.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = Cm(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, top=95, bottom=95)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=font_size, color=WHITE, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = None
        if highlight_rows and ridx in highlight_rows:
            fill = highlight_rows[ridx]
        elif ridx % 2 == 1:
            fill = "F7F9FA"
        for cidx, value in enumerate(row):
            cell = cells[cidx]
            set_cell_margins(cell)
            if fill:
                set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color="27313B", bold=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def keep_table_rows_together(table) -> None:
    """Prevent short comparison tables from leaving a header by itself."""
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=8, color=GRAY)


def add_picture(doc: Document, path: Path, width_inches: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    add_caption(doc, caption)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def create_table5_crop() -> None:
    TABLE5_CROP.parent.mkdir(parents=True, exist_ok=True)
    image = Image.open(PAPER_ASSETS / "table5_recdy.png")
    # Caption, header, and the complete RecDY block on paper page 14.
    crop = image.crop((120, 145, 2160, 1080))
    crop.save(TABLE5_CROP, quality=95)


def build_data() -> dict:
    paper_rows = load_csv(RESULTS / "paper_comparison" / "paper_vs_local_comparison.csv")
    grouped: dict[int, dict[str, dict[str, str]]] = {}
    for row in paper_rows:
        shot = int(row["shot_per_class"])
        grouped.setdefault(shot, {})[row["role"]] = row

    classical = {
        int(row["size"]): row
        for row in load_csv(RESULTS / "improvement" / "classical_improvement_summary.csv")
        if row["condition"] in ("content", "structured_concat", "structured_late_fusion")
    }
    # Preserve all condition rows separately.
    classical_all = load_csv(RESULTS / "improvement" / "classical_improvement_summary.csv")
    roberta = {
        row["tag"]: row
        for row in load_csv(RESULTS / "improvement" / "roberta_improvement_summary.csv")
    }
    gate = load_json(RESULTS / "improvement" / "uncertainty_gate.json")
    return {
        "paper": grouped,
        "classical": classical_all,
        "roberta": roberta,
        "gate": gate,
    }


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for level in (1, 2, 3):
        style = styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(5)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run("提升数据与原论文逐项对比  |  RecDY 推荐意图识别")
    set_run_font(hr, size=8, color=TEAL, bold=True)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bottom = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:bottom")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "8")
    bdr.set(qn("w:space"), "4")
    bdr.set(qn("w:color"), TEAL)
    bottom.append(bdr)
    hp._p.get_or_add_pPr().append(bottom)

    footer = section.footer
    footer.paragraphs[0].text = ""
    add_page_number(footer.paragraphs[0])


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("提升数据与原论文\n逐项对比")
    set_run_font(r, size=25, color=NAVY, bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(18)
    r2 = p2.add_run("RTX 4060 本地实验：性能提升、可比性边界与原文位置索引")
    set_run_font(r2, size=12.5, color=TEAL, bold=True)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.0)
    table.columns[1].width = Cm(12.8)
    metadata = [
        ("对照论文", "Zhu et al., Analyzing bullet chats for recommendation intent identification: Dataset and method, Artificial Intelligence 355 (2026) 104528"),
        ("核心数据集", "RecDY；论文固定测试集，本地严格对照 n = 9,069"),
        ("本地环境", "NVIDIA GeForce RTX 4060 Laptop GPU，8 GB 显存"),
        ("生成日期", "2026-07-18"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        for cell in row.cells:
            set_cell_shading(cell, BLUE_LIGHT)
            set_cell_margins(cell, top=115, bottom=115)
            set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": "C7D4DF"})
        row.cells[0].text = ""
        lr = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(lr, size=8.5, color=NAVY, bold=True)
        row.cells[1].text = ""
        vr = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(vr, size=8.5, color="27313B")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_callout(
        doc,
        "一句话结论",
        "本地低成本方案在内部实验上取得可重复的局部提升；与论文 Table 5 最接近的独立运行仍低 1.23 pp，而 70-shot 三模型工程集成达到 81.65%，高 2.42 pp，但它使用了更大的累计标注覆盖，不能表述为同预算单模型超过 SPT-RII。",
        "success",
    )
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(22)
    r3 = p3.add_run("文档性质：独立实验对比附件  ·  数据、图、位置一一对应")
    set_run_font(r3, size=9, color=GRAY)
    page_break(doc)


def executive_summary(doc: Document) -> None:
    add_heading(doc, "1. 结论先行", 1)
    add_body(
        doc,
        "本项目没有完整复现 SPT-RII 的 soft prompt、BiLSTM prompt interaction、动态 verbalizer 和词性/语义锚更新，而是在 RTX 4060 8 GB 条件下建立了一条低成本替代路线：扩大训练规模、压缩解释、验证集/OOF 融合、三种子概率集成，以及只在低置信样本上调用 Qwen 的不确定性门控。",
    )

    rows = [
        ["严格对齐论文 Table 5", "70-shot 独立运行", "78.00±0.14%", "论文 79.23±0.31%", "-1.23 pp", "最接近同协议"],
        ["工程集成点", "70-shot 三种子概率集成", "81.65%", "论文 79.23±0.31%", "+2.42 pp", "累计标签预算更大"],
        ["内部改进", "TF-IDF，1000 训练", "82.29% → 83.82%", "不同测试协议", "+1.53 pp", "CI 跨 0"],
        ["内部改进", "RoBERTa，1000 训练", "87.19±0.53% → 87.63±0.29%", "不同测试协议", "+0.43 pp", "3 个种子均为正"],
        ["探索性改进", "不确定性门控", "82.29% → 83.71%", "不同测试协议", "+1.42 pp", "95% CI [-2.61,+5.37]"],
    ]
    add_table(
        doc,
        ["比较层级", "实验", "本地结果", "原论文参照", "变化", "判定"],
        rows,
        widths=[2.25, 3.4, 3.65, 3.2, 1.7, 2.35],
        font_size=7.35,
        highlight_rows={0: BLUE_LIGHT, 1: GOLD_LIGHT},
    )
    add_caption(doc, "表 1  本文档中五类结果的层级。只有第一行可作为协议对齐的独立运行对照。")

    add_callout(
        doc,
        "最重要的边界",
        "论文仅写 F1/Prec./Rec.，没有明确 macro、micro 或 weighted averaging；本地采用 scikit-learn Macro-F1/Macro-Precision/Macro-Recall。因此本文给出数值差，但称为“近似口径对照”，不声称指标定义完全等价。",
        "warning",
    )
    add_body(doc, "提升来自什么：", bold_prefix="提升来自什么：")
    add_bullet(doc, "把 400 条训练扩到 1,000 条，首先提高基础分类器稳定性。")
    add_bullet(doc, "不再把冗长解释直接硬拼到输入，而是保留内容通道与解释通道，通过 OOF/验证集选择融合权重和阈值。")
    add_bullet(doc, "RoBERTa 使用更短的 compact explanation；在 1,000 条训练时，三种子差值均为正。")
    add_bullet(doc, "Qwen 不做全量替换，只给低置信样本提供第二意见，减少无效计算与错误覆盖。")
    add_bullet(doc, "三种子概率集成降低单次抽样波动，但扩大了累计训练样本并集，属于工程性能点。")


def location_index(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "2. 原论文位置索引与本地对应关系", 1)
    add_body(doc, "PDF 页码与论文印刷页码一致。下表采用“原论文位置 → 原文作用 → 本地对应 → 可比性”的顺序，读者可直接回到原文核验。")
    rows = [
        ["问题动机", "§5.1，p.8，S0064", "短弹幕依赖上下文；流式语义锚会被稀释", "上下文重构与解释增强", "理论动机"],
        ["解释生成", "§5.2–§5.3，pp.9–10，Fig.5，S0066–S0073", "LLM 在前 20 条滑窗内做语境语义重构", "Qwen 背景、compact explanation", "概念对应"],
        ["提示学习", "§5.4，p.10，S0073/S0077", "Soft prompt + BiLSTM 交互", "未实现；用 RoBERTa/TF-IDF 替代", "未复现"],
        ["动态标签词", "§5.5–§5.6，pp.10–12，S0080–S0101", "POS 候选、语义距离、half-life=10、Top-k 更新", "未实现；验证集融合/门控替代", "未复现"],
        ["数据协议", "§6.1，p.12，S0109", "各平台 70% train / 30% test；测试集只用于最终评估", "RecDY 固定 test n=9,069", "协议对齐"],
        ["训练设置", "§6.3–§6.4，p.13，S0113", "50/60/70-shot；验证集等大；3 次重复；4090", "每类 shot；种子100/101/102；4060", "近似对齐"],
        ["主结果", "§6.5，Table 5，p.14，T001", "RecDY SPT-RII F1 77.84/78.92/79.23", "独立均值与工程集成", "核心对照"],
        ["消融", "§6.6，Fig.6，p.15，S0116/S0117", "去解释 -4.3 pp；去更新 -2.8 pp", "解释通道、late fusion", "概念对应"],
        ["大模型", "§6.7，Table 6，p.16，S0117/S0118", "LLM/PEFT 与 SPT-RII 对照", "Qwen 直接分类/不确定性门控", "协议不同"],
        ["效率", "§6.10，Table 9，p.17，S0186/S0210", "更换解释模型后的 F1 与时间", "Qwen3-8B Q4 生成时间/显存", "描述性对照"],
        ["参数敏感性", "§6.12，Fig.9，pp.19–20，S0264/S0266", "窗口 0→20 提升，过大引入过时噪声", "压缩上下文、验证集融合", "设计依据"],
        ["错误分析", "§6.13，pp.20–21，S0266", "隐式语义、上下文不足等错误", "门控净纠正 5 条；分流诊断", "概念对应"],
    ]
    add_table(
        doc,
        ["主题", "原论文精确位置", "原文内容", "本地对应", "关系"],
        rows,
        widths=[2.0, 4.1, 5.1, 3.5, 1.8],
        font_size=6.8,
        highlight_rows={6: BLUE_LIGHT},
    )
    add_caption(doc, "表 2  原论文与本地实验的位置级映射。S/T 编号来自本项目 paper_reader/source_map。")

    add_callout(
        doc,
        "未覆盖内容",
        "本地没有复现 Table 7 的 BTD 外部数据、Table 8 的跨平台迁移、Table 10 的蒸馏、Table 11/Fig.8 的模板实验，以及 Fig.9 的完整参数扫描。它们只列入索引，不参与提升结论。",
        "info",
    )


def strict_table5(doc: Document, data: dict) -> None:
    page_break(doc)
    add_heading(doc, "3. 主对比：对应原论文 §6.5 / Table 5 / p.14", 1)
    add_callout(
        doc,
        "可比性等级：A-",
        "同为 RecDY 固定测试集 n=9,069，训练/验证按 50、60、70-shot 抽样，并报告 3 次独立运行。差异是模型结构不同、硬件不同，且论文未明确 F1 averaging 与 shot 是否“每类”；本地依据公开代码按每类 shot 执行。",
        "info",
    )
    add_picture(doc, TABLE5_CROP, 6.55, "图 1  原论文 Table 5 的标题、表头与 RecDY 区域，来源：原论文 p.14。")

    paper = data["paper"]
    rows = []
    for shot in (50, 60, 70):
        group = paper[shot]
        ref = group["reported_reference"]
        mean = group["clean_primary_mean"]
        ensemble = group["clean_primary_ensemble"]
        fusion = group["clean_secondary_ensemble"]
        rows.append([
            str(shot),
            f"{percent(float(ref['macro_f1']))} ± {percent(float(ref['sd']))}",
            f"{percent(float(mean['macro_f1']))} ± {percent(float(mean['sd']))}",
            pp(float(mean["delta_vs_paper_pp"])),
            percent(float(ensemble["macro_f1"])),
            pp(float(ensemble["delta_vs_paper_pp"])),
            percent(float(fusion["macro_f1"])),
            pp(float(fusion["delta_vs_paper_pp"])),
        ])
    strict_table = add_table(
        doc,
        ["shot/类", "论文 SPT-RII\nF1±SD", "本地独立均值\nMacro-F1±SD", "独立差值", "内容集成\n单点", "集成差值", "内容+字符\n集成单点", "融合差值"],
        rows,
        widths=[1.25, 2.55, 2.75, 1.65, 2.0, 1.55, 2.15, 1.55],
        font_size=6.85,
        highlight_rows={2: GOLD_LIGHT},
    )
    keep_table_rows_together(strict_table)
    add_caption(doc, "表 3  Table 5 核心 F1 对照。独立运行均值是主结果；两个集成列均为单个工程点，没有 SD。")

    add_body(doc, "严格结论：", bold_prefix="严格结论：")
    add_bullet(doc, "保持每次独立运行的 70-shot/类预算时，本地内容模型为 78.00±0.14%，仍比论文 79.23±0.31% 低 1.23 pp。")
    add_bullet(doc, "三模型内容概率集成达到 81.65%，相对论文报告值高 2.42 pp；再加入字符 TF-IDF 后达到 81.85%，额外仅 +0.20 pp。")
    add_bullet(doc, "集成成员使用不同的训练/验证抽样，不只是同一训练集上的不同初始化，所以集成点不可写成“70-shot 单模型超过论文”。")

    page_break(doc)
    union_rows = [
        ["50", "100", "299", "592", "77.29%"],
        ["60", "120", "360", "713", "78.86%"],
        ["70", "140", "420", "824", "81.65%"],
    ]
    budget_table = add_table(
        doc,
        ["每成员 shot/类", "每成员训练总数", "3 成员训练并集", "训练+验证总并集", "内容集成点"],
        union_rows,
        widths=[3.0, 3.1, 3.2, 3.4, 2.6],
        font_size=7.7,
        header_fill=TEAL,
        highlight_rows={2: GOLD_LIGHT},
    )
    keep_table_rows_together(budget_table)
    add_caption(doc, "表 4  工程集成的实际累计标注覆盖。70-shot 集成整体看到 420 条训练样本，约 210/类。")
    add_callout(
        doc,
        "表述红线",
        "可以说“工程集成点在固定测试集上达到 81.65%，描述性高于论文 Table 5 的 79.23%”；不可以说“我们以同等 70-shot 预算复现并超过 SPT-RII”。",
        "risk",
    )

    add_picture(doc, FIGURES / "fig6_paper_comparison.png", 6.55, "图 2  固定 RecDY 测试集上的论文对齐比较。误差线为独立运行 SD；空心标记为单个集成点。")


def full_metrics(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "4. Table 5 四指标复核", 1)
    add_body(doc, "原论文 Table 5 同时报告 Acc./Prec./Rec./F1。下表把原值与本地独立均值、工程集成单点按 shot 展开，便于逐列核对。")
    rows = [
        ["50", "论文 SPT-RII", "78.41±0.39", "77.63±0.31", "78.96±0.20", "77.84±0.32"],
        ["50", "本地独立均值", "76.59±2.37", "76.11±2.80", "75.74±1.03", "75.52±1.67"],
        ["50", "本地内容集成", "78.21", "77.17", "77.44", "77.29"],
        ["60", "论文 SPT-RII", "79.50±0.75", "78.65±0.49", "79.48±0.39", "78.92±0.56"],
        ["60", "本地独立均值", "78.09±0.42", "77.21±0.31", "77.06±1.44", "77.00±0.89"],
        ["60", "本地内容集成", "79.94", "79.06", "78.69", "78.86"],
        ["70", "论文 SPT-RII", "80.62±0.37", "79.96±0.32", "80.30±0.23", "79.23±0.31"],
        ["70", "本地独立均值", "79.26±0.52", "78.63±0.92", "77.76±0.58", "78.00±0.14"],
        ["70", "本地内容集成", "82.70", "82.13", "81.29", "81.65"],
    ]
    add_table(
        doc,
        ["shot/类", "方法", "Accuracy (%)", "Precision (%)", "Recall (%)", "F1 (%)"],
        rows,
        widths=[1.6, 4.0, 2.7, 2.7, 2.7, 2.7],
        font_size=7.6,
        highlight_rows={6: BLUE_LIGHT, 8: GOLD_LIGHT},
    )
    add_caption(doc, "表 5  原论文 Table 5 与本地结果的四指标复核。集成行是单点。")
    add_callout(
        doc,
        "指标脚注",
        "论文表头为 Prec./Rec./F1，但 §6.4 未写 averaging；本地三项采用宏平均。因此 Accuracy 可直接理解，Precision/Recall/F1 仍需保留口径不完全确定的注记。",
        "warning",
    )
    add_body(doc, "论文内部还有一个需要记录的不一致：Table 5 的 RecDY 50-shot SPT-RII F1 为 77.84±0.32%，Table 6 同一标签写 78.32%。论文未解释差异，所以本报告始终使用 Table 5 作为主参照。")


def internal_improvements(doc: Document, data: dict) -> None:
    page_break(doc)
    add_heading(doc, "5. 内部提升实验：对应 §5.3、§6.6 Fig.6，但协议不同", 1)
    add_callout(
        doc,
        "可比性等级：C",
        "以下实验使用 400 条固定本地测试子集，训练规模为 400 或 1,000；它们用于判断改进是否有效，不能把绝对 F1 直接减去论文 Table 5。",
        "info",
    )
    add_heading(doc, "5.1 字符 TF-IDF：解释分通道 + OOF 晚期融合", 2)
    rows = [
        ["400", "内容基线", "78.88%", "基准", "-"],
        ["400", "旧 Qwen 解释直接拼接", "80.75%", "+1.87 pp", "95% CI [-2.80,+6.69]"],
        ["400", "结构化解释直接拼接", "80.84%", "+1.96 pp", "相对旧提示仅 +0.09 pp"],
        ["400", "旧解释 + OOF 晚期融合", "82.39%", "+3.51 pp", "95% CI [-0.32,+7.35]"],
        ["400", "结构化解释 + OOF 晚期融合", "80.03%", "+1.14 pp", "95% CI [-1.00,+3.41]"],
        ["1000", "内容基线", "82.29%", "基准", "-"],
        ["1000", "结构化解释直接拼接", "82.94%", "+0.65 pp", "-"],
        ["1000", "结构化解释 + OOF 晚期融合", "83.82%", "+1.53 pp", "95% CI [-1.28,+4.34]"],
    ]
    add_table(
        doc,
        ["训练量", "方法", "Macro-F1", "相对同规模内容基线", "不确定性/备注"],
        rows,
        widths=[1.7, 6.0, 2.2, 3.25, 3.65],
        font_size=7.5,
        highlight_rows={3: GREEN_LIGHT, 7: GREEN_LIGHT},
    )
    add_caption(doc, "表 6  TF-IDF 改进实验。置信区间为测试样本配对 bootstrap；跨 0 表示证据尚不足。")
    add_body(doc, "解读：数据量和融合策略确实有价值，但“把解释改成结构化提示”本身没有稳定胜过旧提示。400 条时旧解释 late fusion 最强；1,000 条时结构化 late fusion 相对内容基线提高 1.53 pp。")

    add_heading(doc, "5.2 RoBERTa：compact explanation + 验证集融合", 2)
    rows = [
        ["400", "内容基线", "84.90±1.52%", "基准", "-"],
        ["400", "旧 Qwen 直接拼接", "82.57±0.67%", "-2.34 pp", "退化"],
        ["400", "结构化解释直接拼接", "82.44±1.66%", "-2.46 pp", "退化"],
        ["400", "compact 直接拼接", "84.30±1.75%", "-0.60 pp", "基本追回"],
        ["400", "compact late fusion", "84.88±0.61%", "-0.03 pp", "种子方向不一致"],
        ["1000", "内容基线", "87.19±0.53%", "基准", "-"],
        ["1000", "结构化解释直接拼接", "86.73±1.12%", "-0.46 pp", "仍退化"],
        ["1000", "compact 直接拼接", "87.21±0.51%", "+0.02 pp", "接近持平"],
        ["1000", "compact late fusion", "87.63±0.29%", "+0.43 pp", "3 个种子均为正"],
    ]
    add_table(
        doc,
        ["训练量", "输入/融合", "Macro-F1 均值±SD", "相对内容基线", "判断"],
        rows,
        widths=[1.7, 5.55, 3.1, 3.0, 3.45],
        font_size=7.4,
        highlight_rows={8: GREEN_LIGHT},
    )
    add_caption(doc, "表 7  RoBERTa 三种子结果。1000 条 compact late fusion 的逐种子提升为 +0.154/+0.746/+0.401 pp。")
    add_body(doc, "解读：RoBERTa 对噪声解释很敏感。长解释直接拼接会破坏内容信号；缩短解释并改为 late fusion 后，1,000 条训练下得到小幅且三个种子同方向的提升，但 n=3 仍不足以声称统计显著。")
    add_picture(doc, FIGURES / "fig5_improvement_results.png", 6.55, "图 3  内部改进总览：TF-IDF、RoBERTa 与额外计算开销。误差线定义见原实验记录。")

    page_break(doc)
    add_heading(doc, "5.3 与原论文 Fig.6 的对应方式", 2)
    rows = [
        ["原论文完整 SPT-RII", "78.3%", "-", "同一 SPT-RII 框架"],
        ["原论文 -explanations", "74.0%", "-4.3 pp", "证明解释模块重要"],
        ["原论文 -update", "75.5%", "-2.8 pp", "证明动态更新重要"],
        ["本地 TF-IDF 1000 late fusion", "83.82%", "+1.53 pp vs 本地内容", "不同模型/测试子集"],
        ["本地 RoBERTa 1000 compact late fusion", "87.63±0.29%", "+0.43 pp vs 本地内容", "不同模型/测试子集"],
    ]
    mapping_table = add_table(
        doc,
        ["实验", "F1/Macro-F1", "组件变化", "比较边界"],
        rows,
        widths=[5.7, 3.0, 4.0, 4.1],
        font_size=7.45,
        header_fill=TEAL,
    )
    keep_table_rows_together(mapping_table)
    add_caption(doc, "表 8  原论文 Fig.6 与本地内部消融的概念映射。绝对值不做横向胜负比较。")


def gate_and_efficiency(doc: Document, data: dict) -> None:
    add_heading(doc, "6. 不确定性门控：对应 §6.7 / Table 6，仅概念对照", 1)
    add_body(doc, "原论文 Table 6 比较了 10-shot LLM、1,000 样本 PEFT 与 SPT-RII。本地问题不同：不是让 Qwen 全量替代分类器，而是让它只处理基础模型置信度落在 [0.28, 0.72] 的样本。阈值仅在 200 条验证集上选择，测试集只评估一次。")

    flow = doc.add_table(rows=1, cols=5)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["TF-IDF\n概率", "→", "置信度门控\n[0.28, 0.72]", "→", "仅不确定样本\nQwen 第二意见"]
    for idx, (cell, label) in enumerate(zip(flow.rows[0].cells, labels)):
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        set_run_font(r, size=8.5 if idx % 2 == 0 else 12, color=TEAL if idx % 2 == 0 else GRAY, bold=idx % 2 == 0)
        if idx % 2 == 0:
            set_cell_shading(cell, TEAL_LIGHT)
            set_cell_margins(cell, top=130, bottom=130)
            set_cell_border(cell, top={"val": "single", "sz": 5, "color": TEAL}, bottom={"val": "single", "sz": 5, "color": TEAL})
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    rows = [
        ["TF-IDF 基线", "82.29%", "400", "0%", "基准"],
        ["Qwen 直接分类", "74.70%", "400", "100%", "单独使用更差"],
        ["不确定性门控", "83.71%", "400", "38.5%", "+1.42 pp；95% CI [-2.61,+5.37]"],
    ]
    add_table(
        doc,
        ["方法", "Macro-F1", "测试数", "Qwen 覆盖", "相对基线"],
        rows,
        widths=[4.0, 2.6, 2.3, 2.6, 5.25],
        font_size=7.9,
        highlight_rows={2: GREEN_LIGHT},
    )
    add_caption(doc, "表 9  不确定性门控结果。提升区间跨 0，因此属于探索性证据。")

    add_body(doc, "逐行诊断：154 条被路由样本上，Qwen 纠正 33 个基线错误，同时破坏 28 个原本正确样本，净增 5 条正确预测。混淆矩阵从 [[120,37],[30,213]] 变为 [[125,32],[30,213]]，主要减少了 5 个负类被误报为正类的错误。")
    add_picture(doc, FIGURES / "fig7_uncertainty_gate.png", 6.55, "图 4  不确定性门控：验证集选阈值，固定测试集一次评估。")

    add_callout(
        doc,
        "为什么不能与 Table 6 直接相减",
        "论文 Table 6 的大模型行是另一套 few-shot/PEFT 条件；本地门控只在 400 条子集上评估，还把 TF-IDF 与 Qwen 组合。83.71% 不能写成“比论文 SPT-RII 78.32% 高 5.39 pp”。",
        "risk",
    )

    # The gate figure can fill the preceding page; let Word flow directly into
    # the efficiency section to avoid a trailing blank page from a forced break.
    add_heading(doc, "7. 效率：对应 §6.10 / Table 9 / p.17，描述性对照", 1)
    rows = [
        ["论文", "Llama3.1-8B", "75.79%", "257 ms/条", "RTX 4090", "作为 SPT-RII 解释器"],
        ["论文", "Qwen2.5-7B", "76.74%", "283 ms/条", "RTX 4090", "作为 SPT-RII 解释器"],
        ["论文", "Deepseek-V3", "78.32%", "985 ms/条", "未单列本地卡", "远程/大模型设置"],
        ["论文", "ChatGPT-3.5", "79.09%", "1072 ms/条", "服务端", "远程 API"],
        ["论文", "ChatGPT-4o", "78.60%", "1062 ms/条", "服务端", "远程 API"],
        ["本地", "Qwen3-8B Q4", "解释生成，不是同一 F1", "562 ms/条", "RTX 4060 8GB", "800 条墙钟 450.2 s；峰值 6.40 GiB"],
    ]
    add_table(
        doc,
        ["来源", "模型", "报告 F1", "时间", "硬件", "任务含义"],
        rows,
        widths=[1.5, 3.0, 3.2, 2.6, 2.8, 4.5],
        font_size=7.15,
        highlight_rows={5: GOLD_LIGHT},
    )
    add_caption(doc, "表 10  原论文 Table 9 与本地生成开销。硬件、模型量化、计时范围和任务均不同。")
    add_body(doc, "原论文 Table 9 测的是“更换 SPT-RII 解释生成模型后的整体 F1 和单条时间”；本地 0.562 s/条只统计 Qwen3-8B Q4 的解释生成。Fig.7 还明确排除了 LLM 解释开销。因此这里只能说明 RTX 4060 可以跑通该组件并量化成本，不能做速度优越性结论。")


def limitations_and_next(doc: Document) -> None:
    page_break(doc)
    add_heading(doc, "8. 提升可信度、限制与下一步", 1)
    add_heading(doc, "8.1 目前可以成立的结论", 2)
    add_bullet(doc, "低成本系统已跑通，且所有主结果都有固定测试集、验证集选参或 OOF 选参记录。")
    add_bullet(doc, "70-shot 独立运行均值距离原论文仅 1.23 pp；工程集成点达到 81.65%，但预算边界已单独披露。")
    add_bullet(doc, "1,000 条训练时，TF-IDF late fusion 提高 1.53 pp；RoBERTa compact late fusion 提高 0.43 pp，三种子方向一致。")
    add_bullet(doc, "门控证明“只让大模型处理难例”比全量直接分类更合理，但 400 条测试上的置信区间跨 0。")

    add_heading(doc, "8.2 不能忽略的限制", 2)
    rows = [
        ["指标定义", "论文未明确 F1 averaging；本地用 Macro-F1", "所有跨论文差值标为近似口径"],
        ["方法完整性", "未实现 soft prompt、BiLSTM、动态 verbalizer", "称为低成本替代，不称完整复现"],
        ["集成预算", "不同种子重新抽样，70-shot 三模型训练并集 420", "集成点与独立均值分列"],
        ["同直播间相关性", "12 个 live_id 同时出现在训练/测试", "尚未证明跨直播间泛化"],
        ["上下文跨 split", "解释实验中 81% 训练、99.5% 测试有跨 split 上下文", "无金标签泄漏，但带有转导式同流信息"],
        ["统计功效", "内部测试 n=400；RoBERTa 仅 3 个种子", "不把小幅变化写成显著提升"],
    ]
    add_table(
        doc,
        ["风险", "事实", "文档处理"],
        rows,
        widths=[2.7, 8.0, 6.0],
        font_size=7.6,
        header_fill=RED,
    )
    add_caption(doc, "表 11  结果解释所需的边界条件。")

    add_heading(doc, "8.3 下一轮最值得做的四个实验", 2)
    priorities = [
        ["P1", "把门控扩展到 9,069 条完整固定测试集", "验证 +1.42 pp 是否稳定；报告调用率、时间和 paired bootstrap CI"],
        ["P2", "固定同一 70-shot/类训练集做 5–10 次初始化集成", "把纯模型集成收益与扩大样本并集收益分开"],
        ["P3", "按 live_id 分组划分并重建 split 内上下文", "测真正的跨直播间泛化，消除同流转导因素"],
        ["P4", "系统扫描窗口 k={0,5,10,20} 与解释长度", "直接对应原论文 Fig.9，确定 4060 条件下的最佳上下文预算"],
    ]
    add_table(
        doc,
        ["优先级", "实验", "能解决的问题"],
        priorities,
        widths=[1.6, 7.2, 7.9],
        font_size=7.7,
        header_fill=TEAL,
    )
    add_caption(doc, "表 12  下一轮实验优先级。")

    add_callout(
        doc,
        "最终判断",
        "这次提升是成功的工程进展，但不是已经完成的 SPT-RII 同预算复现。最稳妥的成果表述是：本地替代路线进一步缩小了与论文结果的距离，并给出了一个更高的工程集成点；下一步需要控制累计标签预算和直播间划分，才能把它升级为更强的研究结论。",
        "success",
    )


def evidence_index(doc: Document) -> None:
    # The final callout may naturally spill to the next page.  A second forced
    # break would create a page containing only that callout.
    add_heading(doc, "9. 证据文件索引", 1)
    rows = [
        ["原论文 PDF", r"E:\zotero\storage\7DATFQRR\Zhu 等 - 2026 - Analyzing bullet chats for recommendation intent identification Dataset and method.pdf", "§5–§6、Table 5–10"],
        ["原文位置映射", r"E:\Lab\experiment\paper_reader\source_map.json", "S0064–S0266；T001"],
        ["Table 5 原图", r"E:\Lab\experiment\paper_reader\assets\table5_recdy.png", "原论文 p.14"],
        ["论文对齐清单", r"E:\Lab\experiment\results\paper_comparison\paper_matched_manifest.json", "split、种子、测试零调参"],
        ["Table 5 对比数据", r"E:\Lab\experiment\results\paper_comparison\paper_vs_local_comparison.csv", "论文值、本地均值、集成点"],
        ["独立种子来源", r"E:\Lab\experiment\results\paper_comparison\paper_matched_per_seed_source.csv", "50/60/70-shot 逐种子"],
        ["TF-IDF 改进", r"E:\Lab\experiment\results\improvement\classical_improvement_summary.csv", "400/1000 训练"],
        ["配对区间", r"E:\Lab\experiment\results\improvement\improvement_comparisons.csv", "paired/cluster bootstrap"],
        ["RoBERTa 改进", r"E:\Lab\experiment\results\improvement\roberta_improvement_summary.csv", "三种子均值±SD"],
        ["RoBERTa 配对", r"E:\Lab\experiment\results\improvement\roberta_improvement_paired.csv", "逐种子差值"],
        ["不确定性门控", r"E:\Lab\experiment\results\improvement\uncertainty_gate.json", "阈值、覆盖率、CI"],
        ["门控逐行预测", r"E:\Lab\experiment\results\improvement\uncertainty_gate_test_predictions.csv", "纠正/破坏与混淆矩阵"],
        ["完整实验报告", r"E:\Lab\experiment\report\report.pdf", "方法、全部实验与附录"],
    ]
    for row in rows:
        row[1] = soft_wrap_path(row[1])
    add_table(
        doc,
        ["证据", "绝对路径", "用途"],
        rows,
        widths=[3.0, 10.6, 3.1],
        font_size=6.3,
        header_fill=NAVY,
    )
    add_caption(doc, "表 13  本文档所有数字均可回溯到以上文件。")
    add_body(doc, "文档版本：2026-07-18。生成脚本与本文件同目录，重新运行可从 CSV/JSON 证据生成 Word 和 Markdown。")


def build_markdown(data: dict) -> None:
    paper = data["paper"]
    strict_rows = []
    for shot in (50, 60, 70):
        g = paper[shot]
        ref = g["reported_reference"]
        mean = g["clean_primary_mean"]
        ens = g["clean_primary_ensemble"]
        fusion = g["clean_secondary_ensemble"]
        strict_rows.append(
            f"| {shot} | {percent(float(ref['macro_f1']))} ± {percent(float(ref['sd']))} | "
            f"{percent(float(mean['macro_f1']))} ± {percent(float(mean['sd']))} | "
            f"{pp(float(mean['delta_vs_paper_pp']))} | {percent(float(ens['macro_f1']))} | "
            f"{pp(float(ens['delta_vs_paper_pp']))} | {percent(float(fusion['macro_f1']))} |"
        )
    md = f"""# 提升数据与原论文逐项对比

> RTX 4060 本地实验独立附件。生成日期：2026-07-18。

## 一句话结论

本地与原论文 Table 5 最接近的 70-shot/类独立运行结果为 **78.00±0.14%**，比论文 **79.23±0.31%** 低 **1.23 pp**。三种子内容概率集成达到 **81.65%**，描述性高 **2.42 pp**；但三个成员重新抽样，训练并集为 420 条，因此不能解释为同预算 70-shot 单模型超过论文。

## 比较口径

- 原论文位置：§6.5，Table 5，PDF/印刷 p.14，source_map T001。
- 本地固定测试集：RecDY，n=9,069；种子 100/101/102；测试集零调参。
- 原文仅写 F1/Prec./Rec.，未明确 averaging；本地采用 Macro-F1/Macro-Precision/Macro-Recall。
- 本地未复现 soft prompt、BiLSTM prompt interaction 和动态 verbalizer，是低成本替代路线。

## Table 5 主对比

| shot/类 | 论文 SPT-RII F1±SD | 本地独立 Macro-F1±SD | 差值 | 内容集成单点 | 差值 | 内容+字符集成单点 |
|---:|---:|---:|---:|---:|---:|---:|
{chr(10).join(strict_rows)}

工程集成预算：50/60/70-shot 每成员训练总数为 100/120/140，三个成员训练并集分别为 299/360/420；训练+验证总并集为 592/713/824。

![Table 5 对齐图](../../figures/fig6_paper_comparison.png)

## 内部提升：协议不同，不与 Table 5 直接相减

| 路线 | 基线 | 改进 | 变化 | 说明 |
|---|---:|---:|---:|---|
| TF-IDF，1000 训练 | 82.29% | 83.82% | +1.53 pp | 结构化解释 + OOF late fusion；95% CI [-1.28,+4.34] |
| RoBERTa，1000 训练 | 87.19±0.53% | 87.63±0.29% | +0.43 pp | compact late fusion；三个种子均为正 |
| 不确定性门控 | 82.29% | 83.71% | +1.42 pp | 38.5% 路由；95% CI [-2.61,+5.37] |

原论文对应位置：解释生成见 §5.3（pp.9–10）；消融见 §6.6/Fig.6（p.15）；大模型比较见 §6.7/Table 6（p.16）；效率见 §6.10/Table 9（p.17）。这些是概念对应，不是同协议数值比较。

![内部改进总览](../../figures/fig5_improvement_results.png)

![不确定性门控](../../figures/fig7_uncertainty_gate.png)

## 原论文覆盖矩阵

| 原论文位置 | 本地覆盖 | 结论用途 |
|---|---|---|
| §6.5 / Table 5 / p.14 | 直接协议对齐 | 主对比 |
| §6.6 / Fig.6 / p.15 | 组件思想对应 | 解释/融合动机 |
| §6.7 / Table 6 / p.16 | 概念对应 | 门控，不直接相减 |
| §6.10 / Table 9 / p.17 | 描述性对照 | 4060 可运行与成本 |
| §6.12 / Fig.9 / pp.19–20 | 设计依据 | 压缩上下文 |
| Table 7/8/10/11、Fig.8 | 未覆盖 | 不参与结论 |

## 关键限制

1. 论文没有明确 F1 averaging，也没有在正文明确 shot 是“每类”还是“总数”；本地依据公开代码按每类执行。
2. 三模型工程集成扩大了累计标注覆盖，必须与独立运行均值分开。
3. 12 个 live_id 同时出现在训练/测试；尚未证明跨直播间泛化。
4. 解释实验存在同一直播流的跨 split 上下文，无金标签泄漏，但属于转导式信息。
5. 内部测试仅 400 条，多个置信区间跨 0，不应使用“统计显著”表述。

## 证据入口

- `E:\\Lab\\experiment\\results\\paper_comparison\\paper_vs_local_comparison.csv`
- `E:\\Lab\\experiment\\results\\paper_comparison\\paper_matched_manifest.json`
- `E:\\Lab\\experiment\\results\\improvement\\classical_improvement_summary.csv`
- `E:\\Lab\\experiment\\results\\improvement\\roberta_improvement_summary.csv`
- `E:\\Lab\\experiment\\results\\improvement\\uncertainty_gate.json`
- `E:\\Lab\\experiment\\paper_reader\\source_map.json`
- `E:\\Lab\\experiment\\report\\report.pdf`
"""
    MD_PATH.write_text(md, encoding="utf-8")


def build_docx(data: dict) -> None:
    doc = Document()
    configure_document(doc)
    cover(doc)
    executive_summary(doc)
    location_index(doc)
    strict_table5(doc, data)
    full_metrics(doc)
    internal_improvements(doc, data)
    gate_and_efficiency(doc, data)
    limitations_and_next(doc)
    evidence_index(doc)

    core = doc.core_properties
    core.title = "提升数据与原论文逐项对比"
    core.subject = "RecDY 推荐意图识别：RTX 4060 本地改进与原论文位置映射"
    core.author = "实验项目组"
    core.keywords = "RecDY, SPT-RII, RoBERTa, Qwen, TF-IDF, recommendation intent"
    doc.save(DOCX_PATH)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    create_table5_crop()
    data = build_data()
    build_markdown(data)
    build_docx(data)
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
